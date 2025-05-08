import io

import pandas as pd
from docx import Document
from langchain_core.messages import HumanMessage

from common import file_helper


def xlsx_to_markdown(file_path):
    # Read the Excel file
    df = pd.read_excel(file_path, sheet_name="Sheet1")

    # Convert DataFrame to Markdown
    markdown_table = df.to_markdown(index=False)

    return markdown_table


def docx_to_markdown(file_path):
    # Read the DOCX file
    doc = Document(file_path)
    markdown_content = ""

    for para in doc.paragraphs:
        markdown_content += para.text + "\n\n"

    for table in doc.tables:
        for row in table.rows:
            markdown_content += (
                " | ".join(cell.text.strip() for cell in row.cells) + "\n"
            )
    for i in range(5):
        markdown_content = markdown_content.replace("\n\n", "\n")
    return markdown_content


def markdown_to_bytesio(markdown_content):
    bytes_io = io.BytesIO()
    bytes_io.write(markdown_content.encode("utf-8"))
    bytes_io.seek(0)
    # Reset pointer to the beginning
    return HumanMessage(content=bytes_io.read().decode("utf-8"))


def load_file(file_input):
    if file_input.name.endswith(".pdf"):
        return file_helper.load_pdf_fp(file_input)
    elif file_input.name.endswith(".png"):
        return file_helper.load_image_fp(file_input)
    elif file_input.name.endswith(".txt"):
        return file_helper.load_txt_fp(file_input)
    elif file_input.name.endswith(".xlsx"):
        return markdown_to_bytesio(xlsx_to_markdown(file_input))
    elif file_input.name.endswith(".docx"):
        return markdown_to_bytesio(docx_to_markdown(file_input))

    return None


def load_bytesio(inputs: list[io.BytesIO]):
    return [f for f in [load_file(fp) for fp in inputs] if f is not None]
